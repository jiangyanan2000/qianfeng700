import docx
from docx import Document
from docx.shared import RGBColor, Inches,Pt,Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_UNDERLINE,WD_TAB_ALIGNMENT,WD_TAB_LEADER
import os


doc = Document("12.docx")
# doc.add_heading("nihao",0)
# for para in doc.paragraphs:
#     print(para.text)
s = doc.paragraphs[0]
# s.paragraph_format.space_after = Inches(0.5)
# print(dir(s))
r = s.add_run("胶州市胶东街道办事处文件")
r.font.name = '方正小标宋_GBK'
r._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋_GBK')
r.font.size = Pt(72)
# r.font.bold = True
r.font.color.rgb = RGBColor(255,0,0)


# print(doc.paragraphs[0].text)
for i in range(len(doc.paragraphs)):
    print(str(i),  doc.paragraphs[i].text)

doc.save("12.docx")